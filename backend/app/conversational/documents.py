"""Document upload/list/delete for Conversational AI (no RAG dependency)."""

import asyncio
from pathlib import Path
from typing import List
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, desc
from pydantic import BaseModel

from app.core.config import settings
from app.conversational.database import get_db, User, Document
from app.conversational.auth import get_current_user

router = APIRouter()
UPLOAD_DIR = Path(settings.UPLOAD_DIR)
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


class DocumentResponse(BaseModel):
    id: int
    filename: str
    original_filename: str
    file_type: str
    language: str
    processed: bool
    uploaded_by: int
    created_at: str


async def _extract_pdf(path: str) -> str:
    import fitz
    def run():
        doc = fitz.open(path)
        t = "".join(page.get_text() for page in doc)
        doc.close()
        return t
    return await asyncio.to_thread(run)


async def _extract_docx(path: str) -> str:
    from docx import Document as DocxDocument
    def run():
        doc = DocxDocument(path)
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    return await asyncio.to_thread(run)


async def _extract_txt(path: str) -> str:
    def run():
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    return await asyncio.to_thread(run)


async def _extract_xlsx(path: str) -> str:
    import openpyxl
    def run():
        wb = openpyxl.load_workbook(path)
        parts = []
        for name in wb.sheetnames:
            sheet = wb[name]
            for row in sheet.iter_rows(values_only=True):
                parts.append(" ".join(str(c) for c in row if c is not None))
        return "\n".join(parts)
    return await asyncio.to_thread(run)


async def process_document(file_path: str, file_type: str) -> str:
    ext = file_type.lower()
    if ext == "pdf":
        return await _extract_pdf(file_path)
    if ext == "docx":
        return await _extract_docx(file_path)
    if ext == "txt":
        return await _extract_txt(file_path)
    if ext == "xlsx":
        return await _extract_xlsx(file_path)
    raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")


@router.post("/upload", response_model=DocumentResponse)
async def upload_document(
    file: UploadFile = File(...),
    language: str = "hi",
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    ext = ("." + file.filename.split(".")[-1].lower()) if file.filename else ""
    if ext not in settings.ALLOWED_DOC_EXTENSIONS:
        raise HTTPException(status_code=400, detail="File type not allowed")
    content = await file.read()
    if len(content) > settings.MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=400, detail="File too large")
    if language not in settings.SUPPORTED_LANGUAGES:
        raise HTTPException(status_code=400, detail="Unsupported language")

    temp_path = UPLOAD_DIR / f"conv_{current_user.id}_{file.filename}"
    try:
        temp_path.write_bytes(content)
        extracted = await process_document(str(temp_path), ext.lstrip("."))
        if not extracted.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted")
        file_id = f"{current_user.id}_{file.filename}_{len(extracted)}"
        permanent_path = UPLOAD_DIR / f"{file_id}{ext}"
        temp_path.rename(permanent_path)
        db_doc = Document(
            filename=f"{file_id}{ext}",
            original_filename=file.filename,
            file_path=str(permanent_path),
            file_type=ext.lstrip("."),
            language=language,
            uploaded_by=current_user.id,
            processed=True,
        )
        db.add(db_doc)
        await db.commit()
        await db.refresh(db_doc)
        return DocumentResponse(
            id=db_doc.id,
            filename=db_doc.filename,
            original_filename=db_doc.original_filename,
            file_type=db_doc.file_type,
            language=db_doc.language,
            processed=db_doc.processed,
            uploaded_by=db_doc.uploaded_by,
            created_at=db_doc.created_at.isoformat(),
        )
    except HTTPException:
        raise
    except Exception as e:
        if temp_path.exists():
            temp_path.unlink(missing_ok=True)
        raise HTTPException(status_code=500, detail=f"Document processing failed: {str(e)}")


@router.get("/list", response_model=List[DocumentResponse])
async def list_documents(
    limit: int = 20,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(Document)
        .where(Document.uploaded_by == current_user.id)
        .order_by(desc(Document.created_at))
        .limit(limit)
    )
    docs = r.scalars().all()
    return [
        DocumentResponse(
            id=d.id,
            filename=d.filename,
            original_filename=d.original_filename,
            file_type=d.file_type,
            language=d.language,
            processed=d.processed,
            uploaded_by=d.uploaded_by,
            created_at=d.created_at.isoformat(),
        )
        for d in docs
    ]


@router.delete("/{document_id}")
async def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(Document).where(
            Document.id == document_id,
            Document.uploaded_by == current_user.id,
        )
    )
    doc = r.scalar_one_or_none()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    try:
        p = Path(doc.file_path)
        if p.exists():
            p.unlink()
        await db.delete(doc)
        await db.commit()
        return {"message": "Document deleted successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete failed: {str(e)}")


@router.get("/stats")
async def get_document_stats(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(select(Document).where(Document.uploaded_by == current_user.id))
    docs = r.scalars().all()
    file_types: dict = {}
    languages: dict = {}
    for d in docs:
        file_types[d.file_type] = file_types.get(d.file_type, 0) + 1
        languages[d.language] = languages.get(d.language, 0) + 1
    return {
        "total_documents": len(docs),
        "processed_documents": len([d for d in docs if d.processed]),
        "file_types": file_types,
        "languages": languages,
    }
